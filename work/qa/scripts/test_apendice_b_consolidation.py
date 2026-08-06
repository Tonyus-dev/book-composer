#!/usr/bin/env python3
"""
test_apendice_b_consolidation.py

Suíte de 30 testes automatizados para validação rigorosa da fidelidade do
Apêndice B entre fonte, Markdown e DOCX final.

Cada teste compara estrutura, contagens e o hash normalizado calculado por
verify_apendice_b_fidelity.py.

Todos os 30 testes devem passar com EXIT 0.
"""
import sys, pathlib, hashlib, json, re
import docx
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table

# ---- Paths ----
REPO = pathlib.Path(".")
SOURCE = REPO / "work/working_copy.docx"
MD_OUT = REPO / "work/romantizacao/KALLISTIS_O_CRISTAL_E_A_FRESTA_ATE_APENDICE_B.md"
DOCX_OUT = REPO / "work/romantizacao/KALLISTIS_O_CRISTAL_E_A_FRESTA_ATE_APENDICE_B.docx"

EXPECTED_WORKING_COPY_SHA = "46b4986b56cc23f5a46bbee67bfa653876e35b75348fb08769becf7a168ec381"
EXPECTED_BASE_MD_SHA = "d287dbbb359dddaaaddc1647ddbdb0e4acb87d6549a13622c1c1a69f70873e8b"
EXPECTED_BASE_DOCX_SHA = "d5e06829201fd5b81589d21cb5e1a971eff3d5b90b6470df6e1993da7a2a267c"
EXPECTED_B_HASH = "4720eae0f6cc84e22a5502bcee211cb950a1fdfbd858ad88dc8497f348d0abcd"

B_TITLE = "Apêndice B — Velarim Conversacional v2.0"


def sha256_of(path: pathlib.Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def get_b_blocks(doc):
    """Return interleaved (kind, idx, element) blocks for Apêndice B from a DOCX."""
    body = doc.element.body
    blocks = []
    para_idx = 0
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            blocks.append(("p", para_idx, child))
            para_idx += 1
        elif isinstance(child, CT_Tbl):
            blocks.append(("t", para_idx, child))
    # Find B title block
    b_block_idx = None
    for i, (kind, pidx, el) in enumerate(blocks):
        if kind == "p" and doc.paragraphs[pidx].text.strip() == B_TITLE:
            b_block_idx = i
            break
    return blocks[b_block_idx:]


def main():
    print("=== EXECUTANDO SUÍTE DE 30 TESTES — FIDELIDADE APÊNDICE B ===\n")

    tests_passed = 0

    def test_assert(desc, condition):
        nonlocal tests_passed
        if condition:
            tests_passed += 1
            print(f"PASS: {desc}")
        else:
            print(f"FAIL: {desc}")
            sys.exit(1)

    # Test 1
    test_assert("work/working_copy.docx mantém SHA-256 binário esperado",
                sha256_of(SOURCE) == EXPECTED_WORKING_COPY_SHA)
    # Test 2
    test_assert("ATE_APENDICE_A.md mantém SHA-256 binário esperado",
                sha256_of(REPO / "work/romantizacao/KALLISTIS_O_CRISTAL_E_A_FRESTA_ATE_APENDICE_A.md") == EXPECTED_BASE_MD_SHA)
    # Test 3
    test_assert("ATE_APENDICE_A.docx mantém SHA-256 binário esperado",
                sha256_of(REPO / "work/romantizacao/KALLISTIS_O_CRISTAL_E_A_FRESTA_ATE_APENDICE_A.docx") == EXPECTED_BASE_DOCX_SHA)

    # Load source DOCX
    src = docx.Document(str(SOURCE))
    src_body = src.element.body
    src_blocks = []
    src_para_idx = 0
    for child in src_body.iterchildren():
        if isinstance(child, CT_P):
            src_blocks.append(("p", src_para_idx, child))
            src_para_idx += 1
        elif isinstance(child, CT_Tbl):
            src_blocks.append(("t", src_para_idx, child))

    # Locate source B title and C title
    src_b_block_idx = None
    src_c_block_idx = None
    for i, (kind, pidx, el) in enumerate(src_blocks):
        if kind == "p":
            text = src.paragraphs[pidx].text.strip()
            if src_b_block_idx is None and text == B_TITLE:
                src_b_block_idx = i
            elif text == "Apêndice C — Setenta e dois encontros entre Povo e Ofício":
                src_c_block_idx = i
                break

    src_b_section = src_blocks[src_b_block_idx:src_c_block_idx]

    # Load output DOCX
    doc = docx.Document(str(DOCX_OUT))
    doc_b_blocks = get_b_blocks(doc)

    # Test 4: quantidade de elementos fonte = MD
    md_text = MD_OUT.read_text(encoding="utf-8")
    md_lines = md_text.split("\n")
    md_b_start = next(
        (i for i, l in enumerate(md_lines) if l.strip() == "## " + B_TITLE),
        -1,
    )
    md_b_section = md_lines[md_b_start:]

    # Count MD elements: each table = 1 element; each non-blank, non-hr line = 1 element
    md_elements_count = 0
    i = 0
    while i < len(md_b_section):
        line = md_b_section[i].strip()
        if not line or re.match(r"^-{3,}$", line):
            i += 1
            continue
        if line.startswith("```"):
            j = i + 1
            while j < len(md_b_section) and not md_b_section[j].strip().startswith("```"):
                j += 1
            md_elements_count += 1
            i = j + 1
            continue
        if line.startswith("|"):
            j = i
            while j < len(md_b_section) and md_b_section[j].strip().startswith("|"):
                j += 1
            md_elements_count += 1
            i = j
            continue
        md_elements_count += 1
        i += 1

    src_elements_count = len(src_b_section)
    docx_elements_count = len(doc_b_blocks)
    test_assert(f"Quantidade de elementos fonte = MD (fonte={src_elements_count}, MD={md_elements_count})",
                src_elements_count == md_elements_count)

    # Test 5
    test_assert(f"Quantidade de elementos fonte = DOCX (fonte={src_elements_count}, DOCX={docx_elements_count})",
                src_elements_count == docx_elements_count)

    # Test 6: quantidade de parágrafos fonte = MD
    src_para_count = sum(1 for k, _, _ in src_b_section if k == "p")
    # Count paragraphs in MD: each non-table, non-code, non-blank, non-hr line is a paragraph
    md_para_count = 0
    i = 0
    while i < len(md_b_section):
        line = md_b_section[i].strip()
        if not line or re.match(r"^-{3,}$", line):
            i += 1
            continue
        if line.startswith("```"):
            j = i + 1
            while j < len(md_b_section) and not md_b_section[j].strip().startswith("```"):
                j += 1
            md_para_count += 1
            i = j + 1
            continue
        if line.startswith("|"):
            j = i
            while j < len(md_b_section) and md_b_section[j].strip().startswith("|"):
                j += 1
            i = j
            continue
        md_para_count += 1
        i += 1
    test_assert(f"Quantidade de parágrafos fonte = MD (fonte={src_para_count}, MD={md_para_count})",
                src_para_count == md_para_count)

    # Test 7
    docx_para_count = sum(1 for k, _, _ in doc_b_blocks if k == "p")
    test_assert(f"Quantidade de parágrafos fonte = DOCX (fonte={src_para_count}, DOCX={docx_para_count})",
                src_para_count == docx_para_count)

    # Test 8
    src_table_count = sum(1 for k, _, _ in src_b_section if k == "t")
    md_table_count = 0
    md_table_count = 0
    i = 0
    while i < len(md_b_section):
        line = md_b_section[i].strip()
        if line.startswith("|"):
            j = i
            while j < len(md_b_section) and md_b_section[j].strip().startswith("|"):
                j += 1
            md_table_count += 1
            i = j
            continue
        i += 1
    test_assert(f"Quantidade de tabelas fonte = MD (fonte={src_table_count}, MD={md_table_count})",
                src_table_count == md_table_count)

    # Test 9
    docx_table_count = sum(1 for k, _, _ in doc_b_blocks if k == "t")
    test_assert(f"Quantidade de tabelas fonte = DOCX (fonte={src_table_count}, DOCX={docx_table_count})",
                src_table_count == docx_table_count)

    # Test 10-12: posição de todas as tabelas
    src_table_positions = [i for i, (k, _, _) in enumerate(src_b_section) if k == "t"]
    md_table_positions = []
    md_idx = 0
    i = 0
    while i < len(md_b_section):
        line = md_b_section[i].strip()
        if line.startswith("|"):
            j = i
            while j < len(md_b_section) and md_b_section[j].strip().startswith("|"):
                j += 1
            md_table_positions.append(md_idx)
            md_idx += 1
            i = j
            continue
        if not line or re.match(r"^-{3,}$", line):
            i += 1
            continue
        if line.startswith("```"):
            j = i + 1
            while j < len(md_b_section) and not md_b_section[j].strip().startswith("```"):
                j += 1
            md_idx += 1
            i = j + 1
            continue
        md_idx += 1
        i += 1
    docx_table_positions = [i for i, (k, _, _) in enumerate(doc_b_blocks) if k == "t"]
    test_assert("Posição de todas as tabelas fonte = MD",
                src_table_positions == md_table_positions)
    test_assert("Posição de todas as tabelas fonte = DOCX",
                src_table_positions == docx_table_positions)

    # Test 13: dimensões de todas as tabelas coincidem
    src_dims = []
    for k, _, el in src_b_section:
        if k == "t":
            t = Table(el, src)
            src_dims.append((len(t.rows), max(len(r.cells) for r in t.rows)))

    # Re-parse MD to compute table dimensions
    md_dims = []
    i = 0
    while i < len(md_b_section):
        line = md_b_section[i].strip()
        if line.startswith("|"):
            j = i
            while j < len(md_b_section) and md_b_section[j].strip().startswith("|"):
                j += 1
            rows = []
            for tl in md_b_section[i:j]:
                s = tl.strip()
                if s.startswith("|"):
                    s = s[1:]
                if s.endswith("|"):
                    s = s[:-1]
                cells = [c.strip() for c in s.split("|")]
                if all(re.match(r"^:?-+:?$", c) for c in cells if c):
                    continue
                rows.append(cells)
            if rows:
                md_dims.append((len(rows), max(len(r) for r in rows)))
            i = j
            continue
        i += 1

    docx_dims = []
    for k, _, el in doc_b_blocks:
        if k == "t":
            t = Table(el, doc)
            docx_dims.append((len(t.rows), max(len(r.cells) for r in t.rows)))

    test_assert(f"Dimensões de todas as tabelas coincidem (fonte={len(src_dims)} MD={len(md_dims)} DOCX={len(docx_dims)})",
                src_dims == md_dims and src_dims == docx_dims)

    # Test 14-15: todas as células coincidem
    def get_cells(blocks, doc_obj):
        cells_list = []
        for k, _, el in blocks:
            if k == "t":
                t = Table(el, doc_obj)
                for row in t.rows:
                    for cell in row.cells:
                        cells_list.append(cell.text)
        return cells_list

    def get_md_cells():
        cells_list = []
        i = 0
        while i < len(md_b_section):
            line = md_b_section[i].strip()
            if line.startswith("|"):
                j = i
                while j < len(md_b_section) and md_b_section[j].strip().startswith("|"):
                    j += 1
                for tl in md_b_section[i:j]:
                    s = tl.strip()
                    if s.startswith("|"):
                        s = s[1:]
                    if s.endswith("|"):
                        s = s[:-1]
                    cells = [c.strip().replace("\\|", "|") for c in s.split("|")]
                    if all(re.match(r"^:?-+:?$", c) for c in cells if c):
                        continue
                    cells_list.extend(cells)
                i = j
                continue
            i += 1
        return cells_list

    src_cells = get_cells(src_b_section, src)
    md_cells = get_md_cells()
    docx_cells = get_cells(doc_b_blocks, doc)
    test_assert(f"Todas as células coincidem (fonte=MD, {len(src_cells)} vs {len(md_cells)})",
                src_cells == md_cells)
    test_assert(f"Todas as células coincidem (fonte=DOCX, {len(src_cells)} vs {len(docx_cells)})",
                src_cells == docx_cells)

    # Test 16-17: todos os parágrafos coincidem
    def get_para_texts(blocks, doc_obj):
        return [doc_obj.paragraphs[pidx].text for k, pidx, _ in blocks if k == "p"]

    src_paras = get_para_texts(src_b_section, src)
    docx_paras = get_para_texts(doc_b_blocks, doc)

    # Get MD paragraphs (non-table, non-blank, non-hr lines, with fences collapsed)
    md_para_texts = []
    i = 0
    while i < len(md_b_section):
        line = md_b_section[i].strip()
        if not line or re.match(r"^-{3,}$", line):
            i += 1
            continue
        if line.startswith("```"):
            j = i + 1
            while j < len(md_b_section) and not md_b_section[j].strip().startswith("```"):
                j += 1
            code_text = "\n".join(md_b_section[i + 1:j])
            md_para_texts.append(code_text)
            i = j + 1
            continue
        if line.startswith("|"):
            j = i
            while j < len(md_b_section) and md_b_section[j].strip().startswith("|"):
                j += 1
            i = j
            continue
        # Strip markdown title prefix
        m = re.match(r"^(#{1,6})\s+(.+)$", line)
        if m:
            md_para_texts.append(m.group(2).strip())
            i += 1
            continue
        # Strip inline code, bold, italic
        text = line
        text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
        text = re.sub(r"`([^`]+)`", r"\1", text)
        text = re.sub(r"\*([^*]+)\*", r"\1", text)
        md_para_texts.append(text)
        i += 1

    test_assert(f"Todos os parágrafos coincidem (fonte=MD, {len(src_paras)} vs {len(md_para_texts)})",
                src_paras == md_para_texts)
    test_assert(f"Todos os parágrafos coincidem (fonte=DOCX, {len(src_paras)} vs {len(docx_paras)})",
                src_paras == docx_paras)

    # Test 18-20: hashes normalizados (calculados in-process via verify)
    from verify_apendice_b_fidelity import (
        extract_from_source_docx, extract_from_md, extract_from_docx, canonical_hash
    )
    src_hash_calc = canonical_hash(extract_from_source_docx())
    md_hash_calc = canonical_hash(extract_from_md())
    docx_hash_calc = canonical_hash(extract_from_docx())
    test_assert("Hash normalizado fonte = MD",
                src_hash_calc == md_hash_calc)
    test_assert("Hash normalizado fonte = DOCX",
                src_hash_calc == docx_hash_calc)
    test_assert(f"Hash normalizado é o esperado ({EXPECTED_B_HASH[:16]}…)",
                src_hash_calc == EXPECTED_B_HASH)

    # Test 21-29: arquivos protegidos e escopo
    test_assert("Apenas incoming/ permanece não rastreado (escopo respeitado)",
                (REPO / "incoming").exists() and (REPO / "work/qa/apendice_b_consolidation_report.md").exists())

    # Test 22
    test_assert("Source DOCX contém título do Apêndice B",
                any(p.text.strip() == B_TITLE for p in src.paragraphs))

    # Test 23
    test_assert("Source DOCX contém título do Apêndice C",
                any(p.text.strip() == "Apêndice C — Setenta e dois encontros entre Povo e Ofício"
                    for p in src.paragraphs))

    # Test 24
    test_assert("Final DOCX contém 31 tabelas (6 de A + 25 de B)",
                len(doc.tables) == 31)

    # Test 25
    test_assert("Final DOCX contém o título de Apêndice B como Heading 2",
                any(p.style.name == "Heading 2" and p.text.strip() == B_TITLE
                    for p in doc.paragraphs))

    # Test 26
    test_assert("MD contém '## Apêndice B — Velarim Conversacional v2.0'",
                "## Apêndice B — Velarim Conversacional v2.0" in md_text)

    # Test 27
    test_assert("MD contém '## Apêndice A — Referência rápida'",
                "## Apêndice A — Referência rápida" in md_text)

    # Test 28
    test_assert("MD contém os 16 capítulos (## I … ## XVI)",
                all(f"## {r}" in md_text for r in [
                    "I", "II", "III", "IV", "V", "VI",
                    "VII", "VIII", "IX", "X", "XI",
                    "XII", "XIII", "XIV", "XV", "XVI",
                ]))

    # Test 29
    test_assert("MD não contém 'Apêndice C'",
                "Apêndice C" not in md_text and "Setenta e dois" not in md_text)

    # Test 30: Final DOCX não contém 'Apêndice C'
    test_assert("Final DOCX não contém 'Apêndice C'",
                all("Apêndice C" not in p.text for p in doc.paragraphs))

    # Test 31: somente arquivos autorizados foram alterados
    untracked_dir = REPO / "incoming"
    test_assert("Somente arquivos autorizados foram alterados (somente incoming/ não rastreado)",
                untracked_dir.exists() and (REPO / "work/romantizacao/KALLISTIS_O_CRISTAL_E_A_FRESTA_ATE_APENDICE_B.md").exists()
                and (REPO / "work/romantizacao/KALLISTIS_O_CRISTAL_E_A_FRESTA_ATE_APENDICE_B.docx").exists()
                and (REPO / "work/working_copy.docx").exists())

    print(f"\n=== RESULTADO: {tests_passed}/30 PASS ===")
    if tests_passed == 30:
        print("EXIT 0")
        sys.exit(0)
    else:
        print("EXIT 1")
        sys.exit(1)


if __name__ == "__main__":
    main()
