#!/usr/bin/env python3
"""
build_apendice_b_consolidation.py

Extrai o Apêndice B — Velarim Conversacional v2.0 diretamente de
work/working_copy.docx e o concatena com o conteúdo do compilado
ATE_APENDICE_A, gerando:

  - work/romantizacao/KALLISTIS_O_CRISTAL_E_A_FRESTA_ATE_APENDICE_B.md
  - work/romantizacao/KALLISTIS_O_CRISTAL_E_A_FRESTA_ATE_APENDICE_B.docx

Regras de extração:
  - fonte única: work/working_copy.docx (nunca outro arquivo);
  - inicia em "Apêndice B — Velarim Conversacional v2.0" (inclusive);
  - termina imediatamente antes de "Apêndice C — Setenta e dois..." (exclusive);
  - preserva parágrafos, subtítulos, listas, citações, exemplos, blocos,
    tabelas e células na ordem estrutural real;
  - remove cabeçalhos/rodapés, números de página, sumário, marcas Word,
    comentários e revisões controladas (estes não estão no body XML).

Nenhum arquivo fora do escopo é modificado.
"""
import sys, pathlib, re
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table

REPO = pathlib.Path(".")
SOURCE_DOCX = REPO / "work/working_copy.docx"
BASE_MD = REPO / "work/romantizacao/KALLISTIS_O_CRISTAL_E_A_FRESTA_ATE_APENDICE_A.md"
OUTPUT_MD = REPO / "work/romantizacao/KALLISTIS_O_CRISTAL_E_A_FRESTA_ATE_APENDICE_B.md"
OUTPUT_DOCX = REPO / "work/romantizacao/KALLISTIS_O_CRISTAL_E_A_FRESTA_ATE_APENDICE_B.docx"

B_TITLE = "Apêndice B — Velarim Conversacional v2.0"
C_TITLE = "Apêndice C — Setenta e dois encontros entre Povo e Ofício"


# ------------------------------------------------------------
# Markdown conversion helpers
# ------------------------------------------------------------
def md_escape_inline(text: str) -> str:
    """Escape characters that would break markdown formatting."""
    # Don't escape characters that are inside intentional bold/italic/code spans.
    return text


def paragraph_to_md(p, doc) -> str:
    style = p.style.name
    text = p.text
    stripped = text.strip()

    # No content → skip (let the spacing handle it)
    if not stripped and not text:
        return ""

    # Map styles:
    if style == "Kallistis Chapter":
        # Chapter/appendix title → H2 (same as the other chapters)
        return f"## {stripped}\n"

    if style == "Kallistis LiteraryOpener":
        # Introductory italic paragraph
        # Preserve multi-line as a single italic paragraph
        return f"*{stripped}*\n"

    if style == "Heading 1":
        return f"# {stripped}\n"
    if style == "Heading 2":
        return f"### {stripped}\n"
    if style == "Heading 3":
        return f"#### {stripped}\n"
    if style == "Heading 4":
        return f"##### {stripped}\n"

    if style == "Kallistis Velarim":
        # Multi-line examples — preserve as fenced code block
        if "\n" in text:
            inner = text.rstrip("\n")
            return f"```\n{inner}\n```\n"
        else:
            # Single-line Velarim content (e.g. IPA gloss, single verb)
            return f"`{text}`\n"

    # Default: Kallistis Body Justify or anything else
    return f"{text.rstrip()}\n"


def table_to_md(t: Table) -> str:
    """Convert a docx table to a markdown table.

    Handles merged cells (empty repeated cells) by inserting empty markers
    so each row keeps its declared column count.
    """
    rows = []
    for row in t.rows:
        cell_texts = []
        seen_tcs = set()
        for cell in row.cells:
            tc_id = id(cell._tc)
            if tc_id in seen_tcs:
                # Merged-cell continuation — render empty string
                cell_texts.append("")
                continue
            seen_tcs.add(tc_id)
            text = cell.text.replace("\n", " ").strip()
            # Escape pipe in cell text
            text = text.replace("|", "\\|")
            cell_texts.append(text)
        rows.append(cell_texts)

    if not rows:
        return ""

    n_cols = max(len(r) for r in rows)
    # Pad short rows
    for r in rows:
        while len(r) < n_cols:
            r.append("")

    header = rows[0]
    body = rows[1:] if len(rows) > 1 else []

    lines = []
    lines.append("| " + " | ".join(header) + " |")
    lines.append("| " + " | ".join(["---"] * n_cols) + " |")
    for r in body:
        lines.append("| " + " | ".join(r) + " |")
    return "\n".join(lines) + "\n"


# ------------------------------------------------------------
# Main build logic
# ------------------------------------------------------------
def main():
    if not SOURCE_DOCX.exists():
        print(f"FAIL: fonte não encontrada: {SOURCE_DOCX}")
        sys.exit(1)
    if not BASE_MD.exists():
        print(f"FAIL: compilado-base ausente: {BASE_MD}")
        sys.exit(1)

    doc = Document(str(SOURCE_DOCX))

    # Build interleaved block sequence in body order
    body = doc.element.body
    blocks = []
    para_idx = 0
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            blocks.append(("p", para_idx, child))
            para_idx += 1
        elif isinstance(child, CT_Tbl):
            blocks.append(("t", para_idx, child))

    # Locate B-title and C-title blocks
    b_block_idx = None
    c_block_idx = None
    for i, (kind, pidx, el) in enumerate(blocks):
        if kind == "p":
            text = doc.paragraphs[pidx].text.strip()
            if b_block_idx is None and text == B_TITLE:
                b_block_idx = i
            elif text == C_TITLE:
                c_block_idx = i
                break

    if b_block_idx is None:
        print(f"FAIL: não localizei o título do Apêndice B ({B_TITLE!r})")
        sys.exit(1)
    if c_block_idx is None:
        print(f"FAIL: não localizei o título do Apêndice C ({C_TITLE!r})")
        sys.exit(1)

    print(f"Apêndice B: bloco {b_block_idx} → bloco {c_block_idx - 1}")
    print(f"Total de blocos extraídos: {c_block_idx - b_block_idx}")

    # Convert each block in range to markdown
    md_parts = []
    para_count = 0
    table_count = 0
    cell_count = 0
    row_count = 0
    char_count = 0
    word_count = 0

    for kind, pidx, el in blocks[b_block_idx:c_block_idx]:
        if kind == "p":
            p = doc.paragraphs[pidx]
            md_text = paragraph_to_md(p, doc)
            if md_text:
                md_parts.append(md_text)
                para_count += 1
                char_count += len(p.text)
                word_count += len(re.findall(r"\b\w+\b", p.text, flags=re.UNICODE))
        else:
            t = Table(el, doc)
            md_table = table_to_md(t)
            if md_table:
                # Ensure blank line before and after table for proper MD rendering
                md_parts.append("\n" + md_table + "\n")
                table_count += 1
                row_count += len(t.rows)
                for row in t.rows:
                    for cell in row.cells:
                        cell_count += 1

    appendix_b_md = "\n".join(md_parts).rstrip() + "\n"

    # Combine: base MD (chapters I-XVI + Apêndice A) + Apêndice B
    base_text = BASE_MD.read_text(encoding="utf-8")
    if not base_text.endswith("\n"):
        base_text += "\n"
    # Insert a blank line between Apêndice A end and Apêndice B start
    if not base_text.endswith("\n\n"):
        base_text = base_text.rstrip() + "\n\n"

    output_md = base_text + appendix_b_md

    # Write MD
    OUTPUT_MD.parent.mkdir(parents=True, exist_ok=True)
    OUTPUT_MD.write_text(output_md, encoding="utf-8")
    print(f"Wrote MD: {OUTPUT_MD}")

    # Convert MD to DOCX using the existing converter
    import subprocess
    # Use the same conversion approach as before — inline converter
    convert_md_to_docx(output_md, str(OUTPUT_DOCX))
    print(f"Wrote DOCX: {OUTPUT_DOCX}")

    print(f"\n--- Estatísticas do Apêndice B ---")
    print(f"Parágrafos: {para_count}")
    print(f"Tabelas: {table_count}")
    print(f"Linhas de tabela: {row_count}")
    print(f"Células de tabela: {cell_count}")
    print(f"Caracteres: {char_count}")
    print(f"Palavras: {word_count}")


def convert_md_to_docx(md_text: str, docx_path: str):
    """Convert Markdown text to DOCX using python-docx.

    Mirrors the conversion semantics used by the ATE_APENDICE_A consolidator:
    - Headings (H1-H4), bold, italic, code spans, code blocks, tables
    - Horizontal rules become empty paragraphs
    """
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md_text.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # Blank line
        if not stripped:
            i += 1
            continue

        # Fenced code block
        if stripped.startswith("```"):
            # Find closing fence
            j = i + 1
            while j < len(lines) and not lines[j].rstrip().startswith("```"):
                j += 1
            code_text = "\n".join(lines[i + 1:j])
            p = doc.add_paragraph()
            run = p.add_run(code_text)
            run.font.name = "Consolas"
            run.font.size = Pt(10)
            i = j + 1
            continue

        # Horizontal rule
        if re.match(r"^-{3,}$", stripped):
            doc.add_paragraph("")
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if m:
            level = len(m.group(1))
            title = m.group(2).strip()
            doc.add_heading(title, level=level)
            i += 1
            continue

        # Table
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            j = i
            while j < len(lines) and lines[j].rstrip().startswith("|"):
                table_lines.append(lines[j])
                j += 1

            rows = []
            for tl in table_lines:
                s = tl.strip()
                if s.startswith("|"):
                    s = s[1:]
                if s.endswith("|"):
                    s = s[:-1]
                cells = [c.strip() for c in s.split("|")]
                # Skip separator rows
                if all(re.match(r"^:?-+:?$", c) for c in cells if c):
                    continue
                rows.append(cells)

            if rows:
                n_cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=n_cols)
                table.style = "Light Grid Accent 1"
                for r_idx, row in enumerate(rows):
                    for c_idx, cell_text in enumerate(row):
                        if c_idx < n_cols:
                            cell = table.rows[r_idx].cells[c_idx]
                            cell.text = ""
                            para = cell.paragraphs[0]
                            run = para.add_run(cell_text)
                            if r_idx == 0:
                                run.bold = True
            i = j
            continue

        # Italic-only paragraph (asterisks at start and end)
        if re.match(r"^\*[^*]+\*$", stripped):
            italic_text = stripped[1:-1]
            p = doc.add_paragraph()
            run = p.add_run(italic_text)
            run.italic = True
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        # Regular paragraph with possible bold/italic/code
        p = doc.add_paragraph()
        remaining = stripped
        # Split by ** for bold
        parts = re.split(r"(\*\*[^*]+\*\*)", remaining)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            else:
                # Split by ` for inline code
                sub_parts = re.split(r"(`[^`]+`)", part)
                for sub in sub_parts:
                    if not sub:
                        continue
                    if sub.startswith("`") and sub.endswith("`"):
                        run = p.add_run(sub[1:-1])
                        run.font.name = "Consolas"
                        run.font.size = Pt(10)
                    else:
                        # Split by * for italic
                        italic_parts = re.split(r"(\*[^*]+\*)", sub)
                        for ip in italic_parts:
                            if not ip:
                                continue
                            if ip.startswith("*") and ip.endswith("*") and len(ip) > 2:
                                run = p.add_run(ip[1:-1])
                                run.italic = True
                            else:
                                p.add_run(ip)
        i += 1

    doc.save(docx_path)


if __name__ == "__main__":
    main()
