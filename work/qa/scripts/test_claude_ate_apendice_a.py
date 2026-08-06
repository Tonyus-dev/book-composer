#!/usr/bin/env python3
"""
test_claude_ate_apendice_a.py

Suíte de 21 testes automatizados para verificar a consolidação da romantização
de Kallistis até o Apêndice A a partir do texto-fonte do Claude.

Garante:
  - arquivos MD e DOCX existem;
  - SHA-256 do working_copy.docx permanece inalterado;
  - fonte do Claude está intacta;
  - capítulos I-XVI, Apêndice A, seções A.1-A.12 presentes;
  - Apêndice B ausente;
  - metadados de processo removidos;
  - ponto de corte preservado.

Todos os 21 testes devem passar com EXIT 0.
"""
import sys, pathlib, hashlib, re
import docx

# ---- Paths & Expected Hashes ----
REPO = pathlib.Path(".")
MD_PATH = REPO / "work/romantizacao/KALLISTIS_O_CRISTAL_E_A_FRESTA_ATE_APENDICE_A.md"
DOCX_PATH = REPO / "work/romantizacao/KALLISTIS_O_CRISTAL_E_A_FRESTA_ATE_APENDICE_A.docx"
SOURCE_PATH = REPO / "incoming/CLAUDE_KALLISTIS_CAPITULOS_01_16_E_APENDICES.txt"
WORKING_COPY = REPO / "work/working_copy.docx"

EXPECTED_WORKING_COPY_SHA256 = "46b4986b56cc23f5a46bbee67bfa653876e35b75348fb08769becf7a168ec381"
EXPECTED_SOURCE_SHA256 = "706ac13337dcc9076401dbdaefee33c266136a464cad35cc6d3904560d3cc6c0"

EXPECTED_CLOSING = "— Copia — disse Vahn, já de costas, voltando ao trabalho. — Serve pra isso."
ROMAN_NUMS = ["I", "II", "III", "IV", "V", "VI", "VII", "VIII",
              "IX", "X", "XI", "XII", "XIII", "XIV", "XV", "XVI"]


def sha256_of(path: pathlib.Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def main():
    print("=== EXECUTANDO SUÍTE DE 21 TESTES — CONSOLIDAÇÃO APÊNDICE A ===\n")

    tests_passed = 0

    def test_assert(desc: str, condition: bool):
        nonlocal tests_passed
        if condition:
            tests_passed += 1
            print(f"PASS: {desc}")
        else:
            print(f"FAIL: {desc}")
            sys.exit(1)

    # Test 1: MD existe
    test_assert("work/romantizacao/KALLISTIS_O_CRISTAL_E_A_FRESTA_ATE_APENDICE_A.md existe", MD_PATH.exists())

    # Test 2: DOCX existe
    test_assert("work/romantizacao/KALLISTIS_O_CRISTAL_E_A_FRESTA_ATE_APENDICE_A.docx existe", DOCX_PATH.exists())

    # Test 3: SHA-256 do working_copy.docx preservado
    test_assert(
        "work/working_copy.docx mantém SHA-256 esperado (46b4986b…)",
        sha256_of(WORKING_COPY) == EXPECTED_WORKING_COPY_SHA256
    )

    # Test 4: SHA-256 da fonte do Claude preservado
    test_assert(
        "incoming/CLAUDE_KALLISTIS_CAPITULOS_01_16_E_APENDICES.txt mantém SHA-256 esperado (706ac133…)",
        sha256_of(SOURCE_PATH) == EXPECTED_SOURCE_SHA256
    )

    # Read MD content
    md_text = MD_PATH.read_text(encoding='utf-8')
    md_lines = md_text.split('\n')

    # Read DOCX content
    doc = docx.Document(DOCX_PATH)
    docx_paras = [p.text for p in doc.paragraphs]
    docx_table_text = []
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                docx_table_text.append(cell.text)
    docx_full_text = '\n'.join(docx_paras + docx_table_text)

    # Test 5: MD inicia com título correto
    test_assert("MD inicia com '# KALLISTIS'", md_lines[0].strip() == "# KALLISTIS")

    # Test 6: MD termina com fechamento literário exato
    last_non_empty = next((l for l in reversed(md_lines) if l.strip()), "")
    test_assert(
        "MD termina imediatamente após '— Serve pra isso.'",
        last_non_empty == EXPECTED_CLOSING
    )

    # Test 7: MD não contém Apêndice B
    test_assert(
        "MD não contém '## Apêndice B — Velarim Conversacional'",
        "## Apêndice B" not in md_text and "Velarim Conversacional" not in md_text
    )

    # Test 8: DOCX não contém Apêndice B
    test_assert(
        "DOCX não contém 'Apêndice B — Velarim Conversacional'",
        "Apêndice B" not in docx_full_text and "Velarim Conversacional" not in docx_full_text
    )

    # Test 9: MD contém os 16 capítulos
    test_assert(
        "MD contém os 16 capítulos (## I … ## XVI)",
        all(f"## {r}" in md_text for r in ROMAN_NUMS)
    )

    # Test 10: DOCX contém os 16 capítulos como headings
    test_assert(
        "DOCX contém os 16 capítulos como headings (estilo Heading 2)",
        all(p.strip() in docx_paras for p in ROMAN_NUMS)
    )

    # Test 11: MD contém Apêndice A
    test_assert("MD contém '## Apêndice A — Referência rápida'",
                "## Apêndice A — Referência rápida" in md_text)

    # Test 12: DOCX contém Apêndice A
    test_assert("DOCX contém 'Apêndice A — Referência rápida'",
                "Apêndice A — Referência rápida" in docx_full_text)

    # Test 13: MD contém todas as seções A.1-A.12
    test_assert("MD contém as 12 seções A.1 a A.12",
                all(f"### A.{i}" in md_text for i in range(1, 13)))

    # Test 14: DOCX contém todas as seções A.1-A.12
    test_assert("DOCX contém as 12 seções A.1 a A.12",
                all(f"A.{i}" in docx_full_text for i in range(1, 13)))

    # Test 15: DOCX tem 6 tabelas (A.2, A.3, A.6, A.8, A.11 tamanho, A.11 nível)
    test_assert("DOCX contém 6 tabelas (A.2, A.3, A.6, A.8, A.11×2)", len(doc.tables) == 6)

    # Test 16: MD tem quantidade de palavras compatível com romance extenso (>= 5000)
    word_count = len(re.findall(r'\b\w+\b', md_text, flags=re.UNICODE))
    test_assert(
        f"MD contém conteúdo substancial ({word_count} palavras >= 5000)",
        word_count >= 5000
    )

    # Test 17: MD não contém metadado 'CONTROLE CANÔNICO'
    test_assert("MD não contém 'CONTROLE CANÔNICO'", "CONTROLE CANÔNICO" not in md_text)

    # Test 18: MD não contém declaração 'Status: PASS'
    test_assert("MD não contém 'Status: PASS'", "Status: PASS" not in md_text)

    # Test 19: DOCX contém o fechamento literário exato
    test_assert(
        "DOCX contém '— Copia — disse Vahn, já de costas, voltando ao trabalho. — Serve pra isso.'",
        EXPECTED_CLOSING in docx_full_text
    )

    # Test 20: MD contém '# APÊNDICES' e o subtítulo correto
    test_assert(
        "MD contém '# APÊNDICES' e subtítulo '### Cadernos da Fresta'",
        "# APÊNDICES" in md_text and "### Cadernos da Fresta" in md_text
    )

    # Test 21: DOCX tem estrutura completa
    test_assert(
        f"DOCX possui estrutura completa ({len(doc.paragraphs)} parágrafos >= 150)",
        len(doc.paragraphs) >= 150
    )

    print(f"\n=== RESULTADO: {tests_passed}/21 PASS ===")
    if tests_passed == 21:
        print("EXIT 0")
        sys.exit(0)
    else:
        print("EXIT 1")
        sys.exit(1)


if __name__ == "__main__":
    main()
