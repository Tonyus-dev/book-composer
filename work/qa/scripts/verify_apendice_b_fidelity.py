#!/usr/bin/env python3
"""
verify_apendice_b_fidelity.py

Audita, de forma binária, a fidelidade textual e estrutural do Apêndice B
entre três representações:

  1. work/working_copy.docx                       (fonte)
  2. work/romantizacao/KALLISTIS_O_CRISTAL_E_A_FRESTA_ATE_APENDICE_B.md   (markdown)
  3. work/romantizacao/KALLISTIS_O_CRISTAL_E_A_FRESTA_ATE_APENDICE_B.docx (docx final)

Para cada representação, extrai uma lista canônica de elementos na ordem
real, misturando parágrafos e tabelas:

  [
    {"type": "paragraph", "text": "..."},
    {"type": "table", "rows": [["c1","c2"], ...]},
    ...
  ]

Aplica normalizações permitidas e calcula o SHA-256 normalizado.

Os três hashes normalizados devem ser idênticos. Caso contrário, registra a
primeira divergência exata e termina com EXIT 1.
"""
import sys, pathlib, re, hashlib, json
import docx
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table

REPO = pathlib.Path(".")
SOURCE = REPO / "work/working_copy.docx"
MD_OUT = REPO / "work/romantizacao/KALLISTIS_O_CRISTAL_E_A_FRESTA_ATE_APENDICE_B.md"
DOCX_OUT = REPO / "work/romantizacao/KALLISTIS_O_CRISTAL_E_A_FRESTA_ATE_APENDICE_B.docx"

B_TITLE = "Apêndice B — Velarim Conversacional v2.0"
C_TITLE = "Apêndice C — Setenta e dois encontros entre Povo e Ofício"
B_CONCLUSION = "CONCLUSÃO"


# ============================================================
# Normalization
# ============================================================
def normalize_text(text: str) -> str:
    """Aplica apenas as 6 normalizações permitidas."""
    # 1. CRLF/CR → LF
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # 2. NBSP → espaço comum
    text = text.replace("\u00A0", " ")
    # 3. espaços finais de linha removidos
    lines = text.split("\n")
    lines = [ln.rstrip() for ln in lines]
    text = "\n".join(lines)
    return text


def normalize_cell(cell_text: str) -> str:
    """Célula: aplica normalização textual e une múltiplos parágrafos por \\n."""
    text = normalize_text(cell_text)
    # 4. células com múltiplos parágrafos unidas por \n
    # (python-docx já retorna parágrafos da célula unidos por \n por padrão
    # em cell.text, mas garantimos)
    return text


def normalize_paragraph(text: str) -> str:
    """Parágrafo: aplica normalização textual."""
    return normalize_text(text)


# ============================================================
# Source DOCX extraction (B section only)
# ============================================================
def extract_from_source_docx() -> list:
    doc = Document(str(SOURCE))
    body = doc.element.body
    blocks = []
    para_idx = 0
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            blocks.append(("p", para_idx, child))
            para_idx += 1
        elif isinstance(child, CT_Tbl):
            blocks.append(("t", para_idx, child))

    # Locate B title and C title blocks
    b_block_idx = None
    c_block_idx = None
    for i, (kind, pidx, el) in enumerate(blocks):
        if kind == "p":
            text = doc.paragraphs[pidx].text.strip()
            if b_block_idx is None and text == B_TITLE:
                b_block_idx = i
            elif text == C_TITLE:
                c_block_idx = i
                break

    if b_block_idx is None:
        raise RuntimeError(f"Título do Apêndice B não encontrado em {SOURCE}")
    if c_block_idx is None:
        raise RuntimeError(f"Título do Apêndice C não encontrado em {SOURCE}")

    result = []
    for kind, pidx, el in blocks[b_block_idx:c_block_idx]:
        if kind == "p":
            p = doc.paragraphs[pidx]
            text = normalize_paragraph(p.text)
            result.append({"type": "paragraph", "text": text})
        else:
            t = Table(el, doc)
            rows = []
            for row in t.rows:
                cells = []
                for cell in row.cells:
                    cell_text = cell.text
                    cells.append(normalize_cell(cell_text))
                rows.append(cells)
            result.append({"type": "table", "rows": rows})
    return result


# ============================================================
# MD extraction (B section only)
# ============================================================
def extract_from_md() -> list:
    """Parse markdown between ## Apêndice B — ... and EOF."""
    text = MD_OUT.read_text(encoding="utf-8")
    lines = text.split("\n")

    # Find start (## Apêndice B — Velarim Conversacional v2.0)
    start_idx = None
    for i, line in enumerate(lines):
        if line.strip() == "## " + B_TITLE:
            start_idx = i
            break
    if start_idx is None:
        raise RuntimeError(f"Título do Apêndice B não encontrado em {MD_OUT}")

    # Walk forward, collecting elements
    result = []
    i = start_idx
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # End at next H2 (## ...) that is NOT the B title itself
        # Per prompt: "fim das saídas: final da seção CONCLUSÃO"
        # The MD ends naturally after CONCLUSÃO content. We take everything to EOF.
        # But to be defensive: stop if we hit another ## heading (another chapter)
        if i > start_idx and re.match(r"^##\s+", stripped) and not re.match(r"^###\s+", stripped):
            # This is another H2 (chapter or appendix); stop
            break

        # Fenced code block
        if stripped.startswith("```"):
            # Find closing fence
            j = i + 1
            while j < len(lines) and not lines[j].strip().startswith("```"):
                j += 1
            code_lines = lines[i + 1:j]
            code_text = "\n".join(code_lines)
            result.append({"type": "paragraph", "text": normalize_paragraph(code_text)})
            i = j + 1
            continue

        # Table
        if stripped.startswith("|") and "|" in stripped[1:]:
            table_lines = []
            j = i
            while j < len(lines) and lines[j].strip().startswith("|"):
                table_lines.append(lines[j])
                j += 1

            rows = []
            for tl in table_lines:
                s = tl.strip()
                if s.startswith("|"):
                    s = s[1:]
                if s.endswith("|"):
                    s = s[:-1]
                cells = [c.strip() for c in s.split("|")]
                # Skip separator rows (---|---)
                if all(re.match(r"^:?-+:?$", c) for c in cells if c):
                    continue
                # 6. escapes Markdown de barra vertical revertidos ao texto literal
                cells = [c.replace("\\|", "|") for c in cells]
                # 4. células com múltiplos parágrafos unidas por \n (MD uses single line)
                rows.append(cells)
            result.append({"type": "table", "rows": rows})
            i = j
            continue

        # Heading (### or ####, etc.) — strip title syntax
        m = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if m:
            level = len(m.group(1))
            # Per prompt: "sintaxe Markdown de título removida antes da comparação"
            # Remove the heading prefix; preserve the title text only
            title = m.group(2).strip()
            # But for H2 (## ...) this is a chapter/appendix title — also strip
            if level >= 2:
                result.append({"type": "paragraph", "text": normalize_paragraph(title)})
                i += 1
                continue
            # H1 is the document title
            if level == 1:
                result.append({"type": "paragraph", "text": normalize_paragraph(title)})
                i += 1
                continue

        # Horizontal rule
        if re.match(r"^-{3,}$", stripped):
            i += 1
            continue

        # Blank line
        if not stripped:
            i += 1
            continue

        # Inline-code-only paragraph (just `text`)
        if re.match(r"^`[^`]+`$", stripped):
            inner = stripped[1:-1]
            result.append({"type": "paragraph", "text": normalize_paragraph(inner)})
            i += 1
            continue

        # Italic-only paragraph
        if re.match(r"^\*[^*]+\*$", stripped):
            inner = stripped[1:-1]
            result.append({"type": "paragraph", "text": normalize_paragraph(inner)})
            i += 1
            continue

        # Regular paragraph — strip markdown bold/italic markers
        # Remove `**` for bold and `*` for italic (single * not part of word)
        ptext = stripped
        # Remove bold markers (but keep text)
        ptext = re.sub(r"\*\*([^*]+)\*\*", r"\1", ptext)
        # Remove inline code backticks but keep text
        ptext = re.sub(r"`([^`]+)`", r"\1", ptext)
        # Remove italic markers (single *)
        ptext = re.sub(r"\*([^*]+)\*", r"\1", ptext)
        result.append({"type": "paragraph", "text": normalize_paragraph(ptext)})
        i += 1

    return result


# ============================================================
# DOCX final extraction (B section only)
# ============================================================
def extract_from_docx() -> list:
    doc = Document(str(DOCX_OUT))
    body = doc.element.body
    blocks = []
    para_idx = 0
    for child in body.iterchildren():
        if isinstance(child, CT_P):
            blocks.append(("p", para_idx, child))
            para_idx += 1
        elif isinstance(child, CT_Tbl):
            blocks.append(("t", para_idx, child))

    # Find B title block
    b_block_idx = None
    for i, (kind, pidx, el) in enumerate(blocks):
        if kind == "p":
            text = doc.paragraphs[pidx].text.strip()
            if text == B_TITLE:
                b_block_idx = i
                break
    if b_block_idx is None:
        raise RuntimeError(f"Título do Apêndice B não encontrado em {DOCX_OUT}")

    # Take everything from B title to EOF
    result = []
    for kind, pidx, el in blocks[b_block_idx:]:
        if kind == "p":
            p = doc.paragraphs[pidx]
            text = normalize_paragraph(p.text)
            result.append({"type": "paragraph", "text": text})
        else:
            t = Table(el, doc)
            rows = []
            for row in t.rows:
                cells = []
                for cell in row.cells:
                    cell_text = cell.text
                    cells.append(normalize_cell(cell_text))
                rows.append(cells)
            result.append({"type": "table", "rows": rows})
    return result


# ============================================================
# Canonical serialization
# ============================================================
def canonical_hash(elements: list) -> str:
    """Hash the canonical JSON representation."""
    canonical = json.dumps(
        elements,
        ensure_ascii=False,
        separators=(",", ":"),
        sort_keys=False,
    )
    return hashlib.sha256(canonical.encode("utf-8")).hexdigest()


# ============================================================
# Main
# ============================================================
def main():
    print("=== AUDITORIA DE FIDELIDADE — APÊNDICE B ===\n")

    if not SOURCE.exists():
        print(f"FAIL: fonte ausente: {SOURCE}")
        sys.exit(1)
    if not MD_OUT.exists():
        print(f"FAIL: MD ausente: {MD_OUT}")
        sys.exit(1)
    if not DOCX_OUT.exists():
        print(f"FAIL: DOCX ausente: {DOCX_OUT}")
        sys.exit(1)

    # Extract canonical representations
    print("Extraindo fonte DOCX...")
    src_elements = extract_from_source_docx()
    print(f"  Fonte: {len(src_elements)} elementos")

    print("Extraindo MD...")
    md_elements = extract_from_md()
    print(f"  MD: {len(md_elements)} elementos")

    print("Extraindo DOCX final...")
    docx_elements = extract_from_docx()
    print(f"  DOCX: {len(docx_elements)} elementos\n")

    # Compute hashes
    src_hash = canonical_hash(src_elements)
    md_hash = canonical_hash(md_elements)
    docx_hash = canonical_hash(docx_elements)

    print(f"SOURCE_B_NORMALIZED_SHA256: {src_hash}")
    print(f"MD_B_NORMALIZED_SHA256:     {md_hash}")
    print(f"DOCX_B_NORMALIZED_SHA256:   {docx_hash}\n")

    # Compare hashes
    same_md = (src_hash == md_hash)
    same_docx = (src_hash == docx_hash)
    same_all = same_md and same_docx

    if same_all:
        print("✓ Os três hashes normalizados são IDÊNTICOS.\n")
    else:
        print("✗ Os hashes normalizados NÃO coincidem.\n")
        # Find first divergence
        for i in range(max(len(src_elements), len(md_elements), len(docx_elements))):
            src_e = src_elements[i] if i < len(src_elements) else None
            md_e = md_elements[i] if i < len(md_elements) else None
            docx_e = docx_elements[i] if i < len(docx_elements) else None
            if src_e != md_e or src_e != docx_e:
                print(f"PRIMEIRA DIVERGÊNCIA no índice {i}:")
                print(f"  SOURCE: {json.dumps(src_e, ensure_ascii=False)[:200]}")
                print(f"  MD:     {json.dumps(md_e, ensure_ascii=False)[:200]}")
                print(f"  DOCX:   {json.dumps(docx_e, ensure_ascii=False)[:200]}")
                break

    # Stats
    print("\n=== Estatísticas ===")
    print(f"Total de elementos:")
    print(f"  SOURCE: {len(src_elements)}")
    print(f"  MD:     {len(md_elements)}")
    print(f"  DOCX:   {len(docx_elements)}")
    src_para = sum(1 for e in src_elements if e["type"] == "paragraph")
    md_para = sum(1 for e in md_elements if e["type"] == "paragraph")
    docx_para = sum(1 for e in docx_elements if e["type"] == "paragraph")
    src_tables = [e for e in src_elements if e["type"] == "table"]
    md_tables = [e for e in md_elements if e["type"] == "table"]
    docx_tables = [e for e in docx_elements if e["type"] == "table"]
    print(f"Parágrafos:  SOURCE={src_para}, MD={md_para}, DOCX={docx_para}")
    print(f"Tabelas:     SOURCE={len(src_tables)}, MD={len(md_tables)}, DOCX={len(docx_tables)}")
    if src_tables:
        src_total_cells = sum(sum(len(r) for r in t["rows"]) for t in src_tables)
        src_total_rows = sum(len(t["rows"]) for t in src_tables)
        print(f"Total linhas (fonte): {src_total_rows}")
        print(f"Total células (fonte): {src_total_cells}")

    # Compute and return hashes for in-process use by the test script

    return same_all


if __name__ == "__main__":
    ok = main()
    sys.exit(0 if ok else 1)
