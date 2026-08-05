#!/usr/bin/env python3
"""
audit_velarim.py

Extrai e audita todas as entradas relacionadas ao Velarim do documento-fonte:
regras específicas, tabelas de progressão, estatísticas, mecânicas únicas.

Uso:
  python audit_velarim.py <source_docx> <mechanical_values_json> <output_json>
"""
import sys, json, pathlib, re, hashlib
from docx import Document
from docx.oxml.ns import qn


VELARIM_RE = re.compile(r'\b(Velarim|Vélarim|velarim)\b')
NUMS_RE = re.compile(r'-?\d+(?:[,\.]\d+)?')


def sha256(text):
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def run_text(para):
    return "".join(r.text for r in para.runs if r.text)


def cell_text(cell):
    return "\n".join(p.text for p in cell.paragraphs).strip()


def main():
    if len(sys.argv) != 4:
        print("Uso: audit_velarim.py <source_docx> <mechanical_values_json> <output_json>")
        sys.exit(1)

    src = pathlib.Path(sys.argv[1])
    mv_path = pathlib.Path(sys.argv[2])
    out = pathlib.Path(sys.argv[3])
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document(str(src))
    all_entries = json.loads(mv_path.read_text(encoding="utf-8"))

    # Entries already tagged as velarim in mechanical_values
    mv_velarim = [e for e in all_entries if e.get("categoria") == "velarim"
                  or VELARIM_RE.search(e.get("texto_literal", ""))
                  or VELARIM_RE.search(e.get("nome", ""))]

    # Full scan of paragraphs
    para_hits = []
    current_chapter = ""
    for i, para in enumerate(doc.paragraphs):
        text = run_text(para)
        style = para.style.name if para.style else ""
        if style.startswith("Heading 1") and text.strip():
            current_chapter = text.strip()
        if VELARIM_RE.search(text):
            para_hits.append({
                "para_index": i,
                "chapter": current_chapter,
                "text": text,
                "numbers": NUMS_RE.findall(text),
                "hash": sha256(text),
            })

    # Table scan
    table_hits = []
    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                text = cell_text(cell)
                if VELARIM_RE.search(text):
                    table_hits.append({
                        "table_index": ti,
                        "row": ri,
                        "col": ci,
                        "text": text,
                        "numbers": NUMS_RE.findall(text),
                        "hash": sha256(text),
                    })

    result = {
        "velarim_mechanical_entries": len(mv_velarim),
        "velarim_paragraph_hits": len(para_hits),
        "velarim_table_hits": len(table_hits),
        "paragraph_entries": para_hits,
        "table_entries": table_hits,
        "mechanical_entries": mv_velarim,
    }

    out.write_text(json.dumps(result, indent=2, ensure_ascii=False), encoding="utf-8")
    print(f"[audit_velarim] parágrafos: {len(para_hits)}, células: {len(table_hits)}, mecânicos: {len(mv_velarim)} → {out}")


if __name__ == "__main__":
    main()
